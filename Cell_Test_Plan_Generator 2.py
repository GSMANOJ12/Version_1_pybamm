# -*- coding: utf-8 -*-
"""
Created on Thu Jun 20 09:26:22 2024

@author: INIYANK
"""

import pandas as pd
import numpy as np
import re

#import Cell_Test_Plan_Flowchart

class CellTestPlanGenerator:
    test_plan_status_flag = 0

    def __init__(self) -> None:
       self.declare_variables()

    def declare_variables(self):
        self.doc_command_list = [] #Command Keywords used in the test specification

        self.values_list = ""
        self.registration_values_list = ""
        self.final_values_list = ""

        self.iterate_list = [] #List of items that has to be repeated

        self.test_spec_file_path = '' #Test specification directory
        self.operating_window_file_path = '' #Operating window directory
        self.initialization_parameters_file_path = '' #Hard-coded as of now (later can be intergrated with GUI if needed)

        self.test_procedure_input = ''
        self.cell_type_input = ''
        self.cycler_input = ''
        self.termination_input = 0
        self.test_plan_config_input = 0
        self.registration_parameters_input = []

        self.test_procedure_table_string = ''

        self.temperature_column_string = ''
        self.current_column_string = ''
        self.cycle_count_column_string = ''
        self.soc_column_string = ''
        self.start_soc_column_string = ''
        self.end_soc_column_string = ''
        self.pulse_time_column_string = ''
        self.ch_dch_column_string = ''

        self.all_upper_voltage_limit_values_with_temp = [] #Vdyn,max [V] for all temperature ranges
        self.all_lower_voltage_limit_values_with_temp = [] #Vdyn,min [V] for all temperature ranges

        self.upper_voltage_limit_value = 0
        self.lower_voltage_limit_value = 0

        self.upper_temperature_limit_value = 0 #Global safety limit value for temperature

        self.charge_current_limit_value = 0 #Global safety limit value for temcharge current
        self.discharge_current_limit_value = 0 #Global safety limit value for discharge current

        self.derating_current_limit = 0 #De-rating current limit

        self.all_temperature_values_with_col = []

        self.integrate_operating_window_flag = 0

        self.discharge_current_list = []
        self.charge_current_list = []

        self.row_iterator = 0
        self.column_iterator = 0

        self.source_table_index = 0 #Source table index
        self.ref_table_number = 0 #Reference/Supporting table index

        self.tables = []
        self.tables_list_with_names = []
        self.main_table = []
        self.main_table_expanded = []

        self.base_folder = 'Output'
        self.output_folder = ''

        self.rpt_kapa_flag = 0
        self.charge_nom_flag = 0

        self.cycle_loop_start_flag = 1
        self.cycle_loop_end_flag = 1

        self.temperature_list_with_index = []

        self.loop_start_list = []
        self.loop_end_list = [] 

        self.command_iterator = 0
        self.parameter_iterator = 0
        self.exit_condition_iterator = 0
        self.comment_iterator = 0

    def load_configuration_file(self):
        print('\nLoading Backend Configuration file')

        #Lazy module imports
        import os
        import sys
        import json

        base_dir = 'Configuration File'

        config_path = os.path.join(base_dir, "config.json")

        # Read the config.json file
        try:
            with open(config_path, "r") as file:
                config = json.load(file)
        except FileNotFoundError:
            print(f"Error: config.json not found at {config_path}")
            sys.exit(1)
        
        #Load the configuration values
        self.doc_command_list = config["database"]["doc_command_list"]
        self.cell_cycler_list = config["database"]["cell_cycler_list"]
        #self.output_folder = config["folder_structure"]["output_folder"]

        self.test_procedure_table_string = config["doc_matching_string"]["string_for_test_procedure_tables"]

        #Load supporting table column headings
        self.temperature_column_string = config["supporting_table_keywords"]["temperature_column_string"]
        self.current_column_string = config["supporting_table_keywords"]["current_column_string"]
        self.cycle_count_column_string = config["supporting_table_keywords"]["cycle_count_column_string"]
        self.soc_column_string = config["supporting_table_keywords"]["soc_column_string"]
        self.start_soc_column_string = config["supporting_table_keywords"]["start_soc_column_string"]
        self.end_soc_column_string = config["supporting_table_keywords"]["end_soc_column_string"]
        self.pulse_time_column_string = config["supporting_table_keywords"]["pulse_time_column_string"]
        self.ch_dch_column_string = config["supporting_table_keywords"]["ch_dch_column_string"]

    def clear_all_stored_values(self):
        #Clear the values for the stored values
        self.row_iterator = 0
        self.column_iterator = 0
        
        self.values_list = ''
        self.registration_values_list = ''
        self.final_values_list = ''

        self.tables_list_with_names = []

        self.rpt_kapa_flag = 0
        self.charge_nom_flag = 0

        self.cycle_loop_start_flag = 1
        self.cycle_loop_end_flag = 1

        self.loop_start_list = []
        self.loop_end_list = []

    def execute_main_generate_function(test_spec_input_from_gui, operating_window_input_from_gui, initialization_parameters_file_input_from_gui, test_procedure_input_from_gui, cell_type_input_from_gui, cycler_input_from_gui, termination_input_from_gui, test_plan_config_input_from_gui, registration_input_from_gui):
        self.declare_variables()

        #Load the configuration file and populate the values
        self.load_configuration_file()

        self.process_inputs_from_gui(test_spec_input_from_gui, operating_window_input_from_gui, initialization_parameters_file_input_from_gui, test_procedure_input_from_gui, cell_type_input_from_gui, cycler_input_from_gui, termination_input_from_gui, test_plan_config_input_from_gui, registration_input_from_gui)
        
        #Get the upper and lower voltage limit values
        self.get_parameter_limit_from_opw()

        #Read table information from the document
        self.read_tables_from_word(self.test_spec_file_path)

        self.source_table_index = self.get_source_table_index()

        self.main_table = self.tables_list_with_names[self.source_table_index][1]  
        
        #Get the reference table number
        self.get_reference_table()

        #Remove all the white spaces in the tables (except comment section) - Pre-processing
        self.main_table = self.pre_process_document(self.tables_list_with_names[self.source_table_index][1])

        #Create flowchart
        #Cell_Test_Plan_Flowchart.create_flow_chart(main_table_expanded)

        #Find the column parameter positions
        self.get_column_iterators_positions(len(self.main_table[1]))

        self.create_output_folder()

        #Configure the test plan
        if (self.test_plan_config_input == 1):
            #Find the temperature values
            self.temperature_list_with_index = self.find_temperature_values(self.tables[self.ref_table_number])

            for temp_val_and_index in self.temperature_list_with_index:
                print(f'\nProcess block for -> [{temp_val_and_index[1]} degC]\n')

                #Expand the steps in the main table (expanding the loops)
                self.main_table_expanded = self.expand_the_table_steps(self.main_table, temp_val_and_index[1])

                for val,i in enumerate(self.main_table_expanded):
                    print(f'{val} - {i}')

                #Fill the unknown values
                self.populate_unkown_values(temp_val_and_index[1], temp_val_and_index[0])

                #Integrating operating window
                if (self.integrate_operating_window_flag==1):
                    self.apply_operating_window_limits(self.main_table_expanded)
        
                self.set_registration_paramters()

                self.set_calculation_parameters() 
        
                self.process_main_table_expanded()

                for ii in self.main_table_expanded: 
                   print(ii)
        
                self.set_end()

                self.update_rows_and_columns()

                self.save_output_file(temp_val_and_index[1])

                self.clear_all_stored_values()
        else:
            #Expand the steps in the main table (expanding the loops)
            self.main_table_expanded = self.expand_the_table_steps(self.main_table, r'NA')

            for val,i in enumerate(self.main_table_expanded):
                print(f'{val} - {i}')
        
            #Fill the unknown values
            self.populate_unkown_values(r'NA', r'NA')

            for ii in self.main_table_expanded: 
                   print(ii)

            #Integrating operating window
            if (self.integrate_operating_window_flag==1):
                self.apply_operating_window_limits(self.main_table_expanded)
        
            self.set_registration_paramters()

            self.set_calculation_parameters() 
        
            self.process_main_table_expanded()
        
            self.set_end()

            self.update_rows_and_columns()

            self.save_output_file(r'NA')

            self.clear_all_stored_values()
        
        #Marking the flag for completion of the test plan generation
        ProcessBlockForGUI.test_plan_status_flag = 1

    def procees_inputs_for_pybamm(self,test_spec_input_from_gui,test_procedure_input_from_gui,operating_window_input_from_gui,initialization_parameters_file_input_from_gui,cell_type):
        self.test_spec_file_path = str(test_spec_input_from_gui)
        self.test_procedure_input = str(test_procedure_input_from_gui)
        self.operating_window_file_path = str(operating_window_input_from_gui)    
        self.initialization_parameters_file_path = str(initialization_parameters_file_input_from_gui)
        self.cell_type_input = str(cell_type)
        # self.termination_input = 1
        # self.test_plan_config_input = 1
        # self.registration_parameters_input = ["Aux1"]

    def simulation_function(path,ind,type,ow,initt,cell_type):
        import pybamm

        self.declare_variables()

        self.sim_type=type.lower()
        # print(path,table_index)
        self.load_configuration_file()

        self.procees_inputs_for_pybamm(path,ind,ow,initt,cell_type)

        self.get_parameter_limit_from_opw()

        self.read_tables_from_word(self.test_spec_file_path)

        self.source_table_index = self.get_source_table_index()

        self.main_table = self.tables_list_with_names[self.source_table_index][1]

        self.get_reference_table()
        
        #Remove all the white spaces in the tables (except comment section) - Pre-processing
        self.main_table = self.pre_process_document(self.tables_list_with_names[self.source_table_index][1])

        #Find the column parameter positions
        self.get_column_iterators_positions(len(self.main_table[1]))

        #Configure the test plan
        #Find the temperature values
        if (self.sim_type == "pybamm"):
            #Find the temperature values
            #Expand the steps in the main table (expanding the loops)
            self.main_table_expanded = self.expand_the_table_steps(self.main_table, r'NA')

            #Fill the unknown values
            self.populate_unkown_values( r'NA', r'NA')

            table1=self.main_table_expanded[:]

            #Integrating operating window
            if (self.integrate_operating_window_flag==1):
                self.apply_operating_window_limits(self.main_table_expanded)

        prev_temp=0

        def func_to_val(expr):
            d={"irpt":self.cellcapacity,"inom":self.cellcapacity,"vdyn,max":self.upper_voltage_limit_value,"vdyn,min":self.lower_voltage_limit_value,"tpulse":10,"ipulse":30,"imax,cont":20,}
            dk=list(d.keys())
            exp1=re.findall(r'(\d*\.?\d*)\s*([*/⋅⋅.])\s*([a-zA-Z_][a-zA-Z0-9_]*)|([a-zA-Z_][a-zA-Z0-9_]*)\s*([*/⋅⋅.])\s*(\d*\.?\d*)',expr.replace(" ",""))      #k[0] is unit 
            div,mul=0,0
            flag1=0
            if exp1:
                for k1 in exp1:
                    for k1l in k1:
                        if '/' == k1l:
                            # print("div")
                            div=1
                        elif ('⋅' in k1l or '*' in k1l):
                            # print("mul")
                            mul=1
                        elif k1l.lower() in dk:
                            val1=d[k1l.lower()]
                            # print(val1,k1l.lower(),"ddddddddddddddddddddddd")
                        else:
                            val2r = re.search(r'(\d*\.?\d*)', k1l)
                            val2=val2r.group(0)
                            # print(val2,k1l,")
            else:
                if expr.lower() in dk:
                    val1=d[expr.lower()]
                else:
                    val1 = re.search(r'(\d*\.?\d*)\s*([*/⋅⋅.])\s*([a-zA-Z_][a-zA-Z0-9_]*)|([a-zA-Z_][a-zA-Z0-9_]*)\s*([*/⋅⋅.])\s*(\d*\.?\d*)|',expr.lower())
                    print(val1,"here")
                    val1=val1[0]

            if div and val1:
                fin = (float(val1) / float(val2))
                fin1=fin
                # print(fin1)
            elif mul and val1:
                fin = (float(val1) * float(val2))
                fin1=fin 
                # print(fin1)                           
            else:
                fin1=val1
            # print(fin1,"--------------------------------->")
            return fin1

        def convert_table_to_pybamm(step):
            global prev_temp
            l=list()
            if "command" in step[1].lower():  
                return None
            
            for ik in range(0,len(step)):
                step[ik]=step[ik].replace(" ","")
            # return step

            if step[1].lower()=="settemperature":
                temp = re.findall(r"(-?\d+\.?\d*?)\s*?([A-Za-z°]+)", step[2])                       # st=f"Rest for 5 seconds, temperature={i[2]}"
                temp=temp[0]
                if "C" in temp[1]:
                    temp=temp[0]
                elif "K" in temp[1]:
                    tmp = int(temp[0]) - 273.15
                    temp=tmp
                st="Rest for 20 seconds"
                prev_temp=temp
                return st,temp
                    
            if step[1].lower()=="rest":
                timee=re.findall(r"(\d+\.?\d*?)\s*?([a-zA-Z ]+)",step[3], re.IGNORECASE)                                   #n="".join(n)
                timee=timee[0]
                if "m" in timee[1]:
                    form="minutes"
                elif "h" in timee[1]:
                    form="hours"
                else:
                    form="seconds"
                # else:
                #     print("time format has not valid")
                time=timee[0]+" "+form
                st=f"Rest for {time}"
                return st,prev_temp
            
            rate=""
            # d={"irpt":self.cellcapacity,"inom":self.cellcapacity,"vdyn,max":self.upper_voltage_limit_value,"Vdyn,min":self.lower_voltage_limit_value}
            
            if step[1].lower()=="charge":
                pattern_par = re.finditer(r'(I|V)[<>=]((?:.(?!\bI=|\bV=))+)', step[2])
                results_par = [(m.group(1), m.group(2).strip(',')) for m in pattern_par]
                ct1=0
                if "SOC" in step[2]:
                    d={100: 3.496826, 95: 3.336521, 90: 3.335038, 85: 3.344210, 80: 3.334213, 75: 3.334198, 70: 3.331086, 65: 3.333976, 60: 3.314450, 55: 3.300315, 50: 3.298641, 45: 3.296620, 40: 3.294483, 35: 3.288390, 30: 3.277930, 25: 3.262439, 20: 3.244488, 15: 3.228439, 10: 3.210465, 0: 2.593900}
                    pattern_par = re.findall(r'(soc)[<>=](\d+(\.\d+)?)', step[2].lower())
                    # print(pattern_par)
                var,expr=results_par[0]               #################################################### took only for one experssion because of SOC
                if var.lower()=="i":
                    rate="Charge at "
                    if "C" in expr:
                        rate+=f"{expr}"
                    else:
                        v=func_to_val(expr)
                        rate+=f"{v} A"
                if var.lower()=="v":
                    v=func_to_val(expr)
                    rate="Hold at "
                    rate+=f"{v} V"
                pattern_time=re.findall(r'(t)[<>=](\d+(\.\d+)?|[a-zA-Z_][a-zA-Z0-9_]*)', step[3].lower())
                pattern_ex = re.finditer(r'(I|V)[<>=]((?:.(?!\bI=|\bV=))+)', step[3])
                results_ex = [(m.group(1), m.group(2).strip(',')) for m in pattern_ex]
                if "SOC" in step[3]:
                    d1={100: 3.496826, 95: 3.336521, 90: 3.335038, 85: 3.344210, 80: 3.334213, 75: 3.334198, 70: 3.331086, 65: 3.333976, 60: 3.314450, 55: 3.300315, 50: 3.298641, 45: 3.296620, 40: 3.294483, 35: 3.288390, 30: 3.277930, 25: 3.262439, 20: 3.244488, 15: 3.228439, 10: 3.210465, 0: 2.593900}
                    pattern_par1 = re.findall(r'(soc)[<>=](\d+(\.\d+)?)', step[3].lower())
                    val=d1[int(pattern_par1[0][1])]
                    rate+= f" until {val} V"
                    # print()
                if results_ex:
                    var,expr=results_ex[0]               #################################################### took only for one experssion because of SOC
                    # print(var,expr) 
                    if var.lower()=="i":
                        rate+=" until"
                        if "c" in expr.lower():
                            rate+=f"{expr}"
                        else:
                            v=func_to_val(expr)
                            rate+=f" {v} A"
                    if var.lower()=="v":
                        rate+=" until"
                        v=func_to_val(expr)
                        rate+=f" {v} V" 

                if pattern_time: 
                  var=pattern_time[0][0]
                  expr=pattern_time[0][1]
                  if var.lower()=="t":
                    rate+=" for"
                    v=func_to_val(expr)
                    rate+=f" {v} seconds"          
                return rate,prev_temp


            if step[1].lower()=="discharge":
                pattern_par = re.finditer(r'(I|V)[<>=]((?:.(?!\bI=|\bV=))+)', step[2])
                results_par = [(m.group(1), m.group(2).strip(',')) for m in pattern_par]
                if "SOC" in step[2]:
                    d={100: 3.496826, 95: 3.336521, 90: 3.335038, 85: 3.344210, 80: 3.334213, 75: 3.334198, 70: 3.331086, 65: 3.333976, 60: 3.314450, 55: 3.300315, 50: 3.298641, 45: 3.296620, 40: 3.294483, 35: 3.288390, 30: 3.277930, 25: 3.262439, 20: 3.244488, 15: 3.228439, 10: 3.210465, 0: 2.593900}
                    pattern_par = re.findall(r'(soc)[<>=](\d+(\.\d+)?)', step[2].lower())
                    # print(pattern_par,"lklkmlmlmlmkmkmlkmlkmlm")
                var,expr=results_par[0]               #################################################### took only for one experssion because of SOC
                # print(var,expr)
                
                if var.lower()=="i":
                    rate="Discharge at "
                    if "C" in expr:
                        rate+=f"{expr}"
                    else:
                        v=func_to_val(expr)
                        rate+=f"{v} A"
                        # print("----------------",rate)
                if var.lower()=="v":
                    v=func_to_val(expr)
                    rate="Hold at "
                    rate+=f"{v} V"
                
                pattern_ex = re.finditer(r'(I|V)[<>=]((?:.(?!\bI=|\bV=))+)', step[3])
                results_ex = [(m.group(1), m.group(2).strip(',')) for m in pattern_ex]
                if "SOC" in step[3]:
                    d1={100: 3.496826, 95: 3.336521, 90: 3.335038, 85: 3.344210, 80: 3.334213, 75: 3.334198, 70: 3.331086, 65: 3.333976, 60: 3.314450, 55: 3.300315, 50: 3.298641, 45: 3.296620, 40: 3.294483, 35: 3.288390, 30: 3.277930, 25: 3.262439, 20: 3.244488, 15: 3.228439, 10: 3.210465, 0: 2.593900}
                    pattern_par1 = re.findall(r'(soc)[<>=](\d+(\.\d+)?)', step[3].lower())
                    val=d1[int(pattern_par1[0][1])]
                    rate+= f" until {val} V"
                pattern_time=re.findall(r'(t)[<>=](\d+(\.\d+)?|[a-zA-Z_][a-zA-Z0-9_]*)', step[3].lower())
                              #################################################### took only for one experssion because of SOC
                if results_ex:
                    var,expr=results_ex[0] 
                    if var.lower()=="i":
                        rate+=" until" 
                        if "c" in expr.lower():
                            rate+=f"{expr}"
                        else:
                            v=func_to_val(expr)
                            rate+=f" {v} A"
                    if var.lower()=="v":
                        rate+=" until" 
                        v=func_to_val(expr)
                        rate+=f" {v} V"    
                if pattern_time: 
                  var=pattern_time[0][0]
                  expr=pattern_time[0][1]
                  if var.lower()=="t":
                    rate+=" for"
                    v=func_to_val(expr)
                    rate+=f" {v} seconds"  
      
                return rate,prev_temp
                # for var, expr in results:
                #     if ct1>=1:
                #         rate+=" or "
                #     print(f"{var} = {expr}")
                #     if var.lower()=="i":
                #         if expr.lower() in d:
                #             rate+=d[expr.lower()]+"A"
                #         else:
                #             rate+=expr
                #     if var.lower()=="v":
                #         if expr.lower() in d:
                #             rate+=str(d[expr.lower()])+"V"
                #         else:
                #             rate+=expr
                #     ct1+=1
                # print(rate,"okokkokok")
                # exit_cond=step[4].split("=")
                # print(exit_cond,"llllllllllllll")
                # final=step[1]+" "+rate+" "+exit_cond
            # print(step[1],".........")
                # else:
                #     rate=step[2].split("=")
                #     print(rate)
                # exit_cond=step[3].split("=")
                # print(exit_cond)

        # print("ooooooooooooooo")
        l_step=[]
        # table1.append(['', 'Charge', 'I=Ipulse', 't>tpulse', 'Adjust the SOC to SoCRPT, Set'])
        # table1.append(['', 'Discharge', 'I=Ipulse', 't>tpulse', 'Adjust the SOC to SoCRPT, Set'])
        for ii in table1:
            # print(ii,"------------------------------------------------->")
            st=convert_table_to_pybamm(ii)
            # print(st,ii)
            if "start" in ii[1].lower():
                print("----------",ii[1],'-------------')
            if "end" in ii[1].lower():
                print("----------",ii[1],'-------------')
                l_step.append([str("Discharge at 2 C until 1.9 V"),str(25)+'oC'])
            if st:
               if "hold" in str(st[0]).lower():
                   continue
               print(f'pybamm.step.string("{st[0]}"),',ii)
               l_step.append([str(st[0]),str(st[1])+'oC'])
            # # if st:
            #     print(st[0],st[1])
            #     l_step.append([st[0],st[1]])
            # else:
            #     print(ii,"............>>>>>>>>>>>>>>>>>>>")

        model = pybamm.lithium_ion.DFN()

        param = pybamm.ParameterValues("Mohtat2020")

        print("------------------------------------------------>>>>",self.cellcapacity)

        param.update({
            'Nominal cell capacity [A.h]': self.cellcapacity,
            'Open-circuit voltage at 0% SOC [V]': 2.394,
            'Open-circuit voltage at 100% SOC [V]': 3.428,
            'Lower voltage cut-off [V]': 1.8,
            'Upper voltage cut-off [V]': 3.8,
        })

        # param.process_model(model)
        # geometry = model.default_geometry
        # param.process_geometry(geometry)
        # #
        # mesh = pybamm.Mesh(geometry, model.default_submesh_types, model.default_var_pts,)

        # # # # Step 6: Apply spatial methods
        # disc = pybamm.Discretisation(mesh, model.default_spatial_methods)
        # disc.process_model(model,inplace=False)  # ✅ Now spatial variables arefv ready
        # model.check_well_determined(disc)

        solver = pybamm.CasadiSolver( dt_max=50,atol=1e-6, rtol=1e-6)  #return_solution_if_failed_early=True

        # pybamm.step.string("Rest for 5 seconds", temperature="25oC"),
        #     pybamm.step.string("Rest for 30 minutes"),
        #     pybamm.step.string("Discharge at 1.0 A until 2.0 V"),
        #     pybamm.step.string("Discharge at 2.0 V until 1.2 A"),
        #     pybamm.step.string("Rest for 30 minutes"),
        #     pybamm.step.string("Charge at 1.0 A until 3.8 V")
        experiment1 = pybamm.Experiment([
            pybamm.step.string(i[0],temperature=i[1]) for i in l_step
            ])

        sim1 = pybamm.Simulation(model, parameter_values=param,experiment=experiment1,solver=solver)
        solution = sim1.solve([0,80000])
        sim1.plot(["Terminal voltage [V]","Current [A]"])
        # # solution = sim.solve([0, 3600])
        # plot = pybamm.QuickPlot(sim1, figsize=(14, 7))
        # plot.plot(0.5)
        # import matplotlib.pyplot as plt
        
        # # # Extract variables from solution
        # t = solution["Time [s]"].entries
        # voltage = solution["Terminal voltage [V]"].entries
        # current = solution["Current [A]"].entries
        
        # # # Plot manually
        # plt.figure()
        # plt.plot(t, voltage, label="Voltage (V)")
        # plt.plot(t, current, label="Current (A)")
        # plt.xlabel("Time (s)")
        # plt.ylabel("Value")
        # plt.legend()
        # plt.show()
        # def get_specific_capacity_positive(solution):
            # current = solution["Current [A]"].entries
            # time = solution["Time [s]"].entries
            # dt = np.diff(time)
            # current_avg = 0.5 * (current[:-1] + current[1:])
            # capacity = np.abs(np.sum(current_avg * dt)) / 3600  # Convert to Ah
            # return capacity
        # def run_cycle_simulation(model, parameter_values, num_cycles):
                # capacities = []
                # # Define one charge + discharge cycle
                # cycle = [
                #     "Discharge at 1C until 2.5V",
                #     "Charge at 1C until 4.2V"
                # ]
                # experiment_steps = cycle * num_cycles
                # experiment = pybamm.Experiment(experiment_steps)
                # sim = pybamm.Simulation(model, parameter_values=parameter_values, experiment=experiment)
                # solution = sim.solve()
                # # Get step-wise solutions directly
                # split_solutions = solution.cycles
                # for i in range(0, len(split_solutions), 2):  # Only Discharge steps
                #     discharge_solution = split_solutions[i]
                #     capacity = get_specific_capacity_positive(discharge_solution)
                #     capacities.append(capacity)
                # return capacities
        # model = pybamm.lithium_ion.DFN()  # You can also use pybamm.lithium_ion.Chen2020()
        # param = model.default_parameter_values
        # num_cycles = 5

        # capacities = run_cycle_simulation(model, param, num_cycles)

        # # Plotting
        # plt.figure(figsize=(8, 5))
        # plt.plot(range(1, num_cycles + 1), capacities, marker='o')
        # plt.xlabel("Cycle Number")
        # plt.ylabel("Specific Capacity (Positive) [Ah]")
        # plt.title("Specific Capacity (Positive) vs Number of Cycles")
        # plt.grid(True)
        # plt.tight_layout()
        # plt.show()

    def process_inputs_from_gui(self, test_spec_input_from_gui, operating_window_input_from_gui, initialization_parameters_file_input_from_gui, test_procedure_input_from_gui, cell_type_input_from_gui, cycler_input_from_gui, termination_input_from_gui, test_plan_config_input_from_gui, registration_input_from_gui):
        self.test_spec_file_path = str(test_spec_input_from_gui)

        self.operating_window_file_path = str(operating_window_input_from_gui)
        
        self.initialization_parameters_file_path = str(initialization_parameters_file_input_from_gui)

        self.test_procedure_input = str(test_procedure_input_from_gui)

        self.cell_type_input = str(cell_type_input_from_gui)
        
        self.cycler_input = str(cycler_input_from_gui)
        
        self.termination_input = int(termination_input_from_gui)

        self.test_plan_config_input = int(test_plan_config_input_from_gui)

        self.registration_parameters_input = registration_input_from_gui

    def pre_process_document(self, x_list):
        print('\nPreprocessing the table\n')
        x1 = []
        for i in range (0,len(x_list)):
            temp1= []
            for j in range (0,len(x_list[0])-1):
                temp = x_list[i][j][:]

                #Remove the whitespaces in the document
                temp = temp.replace(" ", "")

                #Replace the "⋅" character (To eliminate the encoding issue) #Comment if the character is replaced in t.s
                temp = temp.replace("⋅",".")
                
                #Combine the multiple lines to single line and remove white lines
                temp = temp.strip()
                
                #Replace the escape character(\n) with comma
                temp = re.sub(r'[\n\t\r]', ',', temp)
                temp1.append(temp)
                
            val = x_list[i][len(x_list[0]) -1][:]
            temp1.append(val)
            
            x1.append(temp1[:])
        return x1

    def get_source_table_index(self):
        index = 0

        #Check for the table with the selected input from GUI
        for val,table in enumerate(self.tables_list_with_names):
            if (self.test_procedure_input in table[0]):
                index = val
        return index

    def get_reference_table(self):
        table_index = 0
        reference_table_number_string = ''
        table_pattern = r'Table\s*(\d+\*?)' #Defining the pattern for Table x or Table x*

        #Find the table number from the main table
        for i in range(0,len(self.main_table)):
            if (re.search(table_pattern, self.main_table[i][4])):
                reference_table_number_string = re.findall(table_pattern, self.main_table[i][4])[0]

        #Find the corresponding table index value
        for val,index in enumerate(self.tables_list_with_names):
            if (reference_table_number_string in index[0]):
                table_index = val

        self.ref_table_number = table_index

    def get_column_iterators_positions(self, columns):        
        #Finding the Column Iterator values
        j_index = 0
        for j in range (0,columns):
            if (self.main_table[0][j] == 'Command'):
                self.command_iterator = j_index
            elif (self.main_table[0][j] == 'Parameter'):
                self.parameter_iterator = j_index
            elif (self.main_table[0][j] == 'ExitCondition'):
                self.exit_condition_iterator = j_index
            elif (self.main_table[0][j] == 'Comment'):
                self.comment_iterator = j_index
            j_index += 1

    def read_tables_from_word(self, file_path):
        #Lazy module imports
        from docx import Document

        # Load the document
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        paragraphs_list = list(paragraphs)
        
        tables_data = []
        tables_name=[]
        tables_list_for_gui = [] #list of tables for GUI

        for table in doc.tables:
            #print(table.value)
            # Extract table data
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables_data.append(table_data)

            self.tables = tables_data

            # Extract table name
            for line_index, paragraph in enumerate(paragraphs_list):
                preceding_paragraph = None
                #Find the paragraph one line above the table
                if line_index+1 < len(paragraphs_list) and paragraphs_list[line_index+1]._element.getnext() is table._element:
                    preceding_paragraph = paragraph.text.strip()    
                    tables_name.append(preceding_paragraph)
                    if (self.test_procedure_table_string in preceding_paragraph.lower()):
                        tables_list_for_gui.append(preceding_paragraph)
                    break

        # Combining the table data with table name
        for val in range(0, len(tables_name)):
            self.tables_list_with_names.append([tables_name[val], tables_data[val]])
        
    def get_parameter_limit_from_opw(self):
        #Lazy module imports
        import math
        
        temperature_string = 'Temperature[°C]'.lower()    
        upper_voltage_string = 'Vdyn,max[V]'.lower()
        lower_voltage_string = 'Vdyn,min[V]'.lower()
        capac="Cellcapacity[Ah]:".lower()

        all_temperature_values = []
        all_upper_voltage_limit_values = []
        all_lower_voltage_limit_values = []
        capacity_value=0
        flag_cap=0

        #Load the excel file
        sheet_data = pd.read_excel(self.operating_window_file_path, sheet_name=self.cell_type_input)

        #Fetching the values of the temperature
        for row_index, row in sheet_data.iterrows():
            for col_index, value in enumerate(row):
                #Check for 'Temperature[°C]' string in the excel
                if flag_cap==1 and str(value) != "nan":
                    # print(value)
                    if value:
                        capacity_value=value
                        flag_cap=0

                if (pd.notna(value) and temperature_string in str(value).lower().replace(" ","")):
                    #Store all the numerical values from the particular row and removes NaN
                    all_temperature_values = [value for value in row if isinstance(value, (int, float)) and not math.isnan(value)]
                    #Store all temperature values with column index
                    self.all_temperature_values_with_col = [[value,index] for index,value in enumerate(row) if isinstance(value, (int, float)) and not math.isnan(value)]

                #Check for 'Vdyn,max[V]' string in the excel
                elif (pd.notna(value) and upper_voltage_string in str(value).lower().replace(" ","")):
                    #Store all the numerical values from the particular row and removes NaN
                    all_upper_voltage_limit_values = [value for value in row if isinstance(value, (int, float)) and not math.isnan(value)]

                #Check for 'Vdyn,min[V]' string in the excel
                elif (pd.notna(value) and lower_voltage_string in str(value).lower().replace(" ","")):
                    #Store all the numerical values from the particular row and removes NaN
                    all_lower_voltage_limit_values = [value for value in row if isinstance(value, (int, float)) and not math.isnan(value)]

                elif (pd.notna(value) and capac in str(value).lower().replace(" ","")):
                    #Store all the numerical values from the particular row and removes NaN
                    flag_cap = 1
                
        self.cellcapacity=capacity_value

        for i in range(0, len(all_upper_voltage_limit_values)):
            #Appending the upper voltage limit with respective to the temperature
            self.all_upper_voltage_limit_values_with_temp.append(list([all_temperature_values[i], all_upper_voltage_limit_values[i]]))

            #Appending the lower voltage limit with respective to the temperature
            self.all_lower_voltage_limit_values_with_temp.append(list([all_temperature_values[i], all_lower_voltage_limit_values[i]]))

        #Slicing the list and removing the duplicate entries for temperature
        self.all_temperature_values_with_col = self.all_temperature_values_with_col[:len(self.all_upper_voltage_limit_values_with_temp)]

        charge_current_string = 'Itherm,cont,ch[A]'.lower()
        discharge_current_string = 'Itherm,cont,dch[A]'.lower()

        discharge_plating_current_string = 'Iplat,cont[A]'.lower()

        charge_current_values = []
        discharge_current_values = []
        discharge_plating_current_values = []
        
        #Fetching the values for the current
        for row_index, row in sheet_data.iterrows():
            for col_index, value in enumerate(row):
                if (pd.notna(value) and charge_current_string in str(value).lower().replace(" ","")):
                    for temp_value in self.all_temperature_values_with_col:
                        temperature_col_index = temp_value[1]
                        
                        charge_loop_iterator_start_flag = 0
                        charge_loop_iterator_end_flag = 0

                        for val1, val2 in zip(sheet_data.iloc[row_index:,temperature_col_index], sheet_data.iloc[row_index:,col_index-1]):
                            #Charge loop start condition
                            if ((val2 == 0) and charge_loop_iterator_end_flag == 0):
                                charge_loop_iterator_start_flag= 1

                            #Fetching the charge steps from operating window
                            if (charge_loop_iterator_start_flag==1 and charge_loop_iterator_end_flag==0):
                                charge_current_values.append(val1)

                            #Charge loop end condition
                            if ((val2==100) and charge_loop_iterator_start_flag == 1):
                                charge_loop_iterator_end_flag = -1
                
                elif (pd.notna(value) and discharge_current_string in str(value).lower().replace(" ","")):
                    for temp_value in self.all_temperature_values_with_col:
                        temperature_col_index = temp_value[1]
                        
                        discharge_loop_iterator_start_flag = 0
                        discharge_loop_iterator_end_flag = 0

                        for val3, val4 in zip(sheet_data.iloc[row_index:,temperature_col_index], sheet_data.iloc[row_index:,col_index-1]):
                            #Fetching the charge steps from operating window
                            if (discharge_loop_iterator_start_flag==1 and discharge_loop_iterator_end_flag==0):
                                discharge_current_values.append(val3)
                            
                            #Discharge loop start condition
                            if ((val3 == 0) and discharge_loop_iterator_end_flag == 0):
                                discharge_loop_iterator_start_flag= 1
                            #Discharge loop end condition
                            if ((pd.isna(val4) or val4==100) and discharge_loop_iterator_start_flag == 1):
                                discharge_loop_iterator_end_flag = -1

                elif (pd.notna(value) and discharge_plating_current_string in str(value).lower().replace(" ","")):
                    for temp_value in self.all_temperature_values_with_col:
                        if (temp_value[0] == 25):
                            temperature_col_index = temp_value[1]
                            break
                        
                    discharge_plating_loop_iterator_start_flag = 0
                    discharge_plating_loop_iterator_end_flag = 0

                    for val5, val6 in zip(sheet_data.iloc[row_index:,temperature_col_index], sheet_data.iloc[row_index:,col_index-1]):
                        #print(val5, val6)
                        #Fetching the charge steps from operating window
                        if (discharge_plating_loop_iterator_start_flag==1 and discharge_plating_loop_iterator_end_flag==0):
                            discharge_plating_current_values.append(val5)
                        
                        #Discharge loop start condition
                        if ((val6 == 0) and discharge_plating_loop_iterator_end_flag == 0):
                            discharge_plating_loop_iterator_start_flag= 1
                        #Discharge loop end condition
                        if ((val6==100) and discharge_plating_loop_iterator_start_flag == 1):
                            discharge_plating_loop_iterator_end_flag = -1
        
        #Fetch the parameters for the global safety limits
        #Upper and lower voltage limits (first value) - need to change for dynamic temperature ranges
        self.upper_voltage_limit_value = self.all_upper_voltage_limit_values_with_temp[0][1]  #vmax
        self.lower_voltage_limit_value = self.all_lower_voltage_limit_values_with_temp[0][1]  #vmin

        #Upper temperature limit
        self.upper_temperature_limit_value = self.all_temperature_values_with_col[-1][0]

        #Charge current limit
        self.charge_current_limit_value = max(charge_current_values)

        #Discharge current limit
        self.discharge_current_limit_value = max(discharge_current_values)

        #De-rating current limit
        self.derating_current_limit = discharge_plating_current_values[-1]

    def get_registration_parameters(file_path):
        registration_parameters_for_gui = []
        
        reg_sheet_name = "Registration Format"
        ini_sheet_name = "Additional Parameters"
        sheet_1_column_header = "Parameter"
        sheet_2_column_header = "Variable"

        #Load the excel file and process registration parameters
        registration_sheet_data = pd.read_excel(file_path, sheet_name=reg_sheet_name)
        initialization_sheet_data = pd.read_excel(file_path, sheet_name=ini_sheet_name)

        if sheet_1_column_header in registration_sheet_data.columns:
            for value_1 in registration_sheet_data[sheet_1_column_header]:
                registration_parameters_for_gui.append(str(value_1))

        if sheet_2_column_header in initialization_sheet_data.columns:
            for value_2 in initialization_sheet_data[sheet_2_column_header]:
                registration_parameters_for_gui.append(str(value_2))    
            
        return registration_parameters_for_gui

    def find_cycle_positions(self, rows):
        start_flag = ''
        
        for i in range(0,rows):
            #Find the position of Cycle-start
            if ('CYCLE-START' in self.main_table[i][self.command_iterator]):
                self.loop_start_list.append(i)
                start_flag = self.main_table[i][self.command_iterator][-1]
            #Find the position of Cycle-end
            for j in range(0,rows):
                if ('CYCLE-END' in self.main_table[j][self.command_iterator] and self.main_table[j][self.command_iterator][-1] == start_flag):
                    self.loop_end_list.append(j)
                    start_flag = ''
                    break

    def find_temperature_values(self, input_file):
        temp_list = []
        temp_list_with_temp = []
        ref_temperature_column = 0

        for col_val in range(0,len(input_file[0])):
            #Keyword search for the cycle count values in the table
            if (input_file[0][col_val].replace(" ", "").lower() == self.temperature_column_string):
                ref_temperature_column = col_val

        temp_list.append(input_file[1][ref_temperature_column])
        temp_list_with_temp.append([1, input_file[1][ref_temperature_column]])

        for index, row_val in enumerate(input_file[2:]):
            if (row_val[ref_temperature_column] != temp_list[-1]):
                temp_list.append(row_val[ref_temperature_column])
                temp_list_with_temp.append([index+2, row_val[ref_temperature_column]])

        return temp_list_with_temp

    def add_flag_bit(self, rows):
        inner_expand_start = 0
        inner_expand_end = 0

        #Initializing the loop position bit to -1
        loop_position_bit = -1

        #Defining the checklist
        parameter_check_list_keywords = []
        exit_condition_check_list_keywords = ["Ah-Set<-XCRPT,I<0.05.INOM", "Ah-Set>-XCRPT,I<0.05.INOM"]
        
        for i in range(0,rows):
            #Update the cycle-list position
            if ('CYCLE-START' in self.main_table[i][self.command_iterator]):
                loop_position_bit += 1
            
            #Configuring the keywords for searching in the loop
            parameter_check_list_keywords = ["I=xx"] if self.test_plan_config_input == 1 else ["I=xx", "T=TSet", "T=xx"] 

            #Iniyan



            if (self.main_table[i][self.parameter_iterator] in parameter_check_list_keywords):
                #Initializing the duplicate flag
                duplicate_flag_bit = 0
                
                #Finding the Cycle start and end positions
                inner_expand_start = self.loop_start_list[loop_position_bit]
                inner_expand_end = self.loop_end_list[loop_position_bit]

                for j in range(self.loop_start_list[loop_position_bit], self.loop_end_list[loop_position_bit]+1):
                    #Check for duplicate entries of flag bit
                    if (len(self.main_table[j]) == len(self.main_table[0])+1):
                        duplicate_flag_bit = 1
                    
                    #Append flag bit only for the non-duplicate entries
                    if (duplicate_flag_bit == 0):
                        if ('CYCLE-START' in self.main_table[j][self.command_iterator]):
                            #Appending '2' at the last for Cycle-start of iterating loop
                            self.main_table[inner_expand_start].append('2')
                        elif ('CYCLE-END' in self.main_table[j][self.command_iterator]):
                            #Appending '-2' at the last for Cycle-start of iterating loop
                            self.main_table[inner_expand_end].append('-2')
                        else:
                            #Appending '1' at the last for elements of iterating loop
                            self.main_table[j].append('1')

    def find_iterate_list(self, start_row_val, rows):
        list_append_flag = 0
        self.iterate_list = []

        cycle_start_position = start_row_val
        cycle_end_position = 0


        print('iterate list----------------------')

        print(self.loop_start_list)
        print(self.loop_end_list)

        for index, value in enumerate(self.loop_start_list):
            if (value == start_row_val):
                cycle_end_position = self.loop_end_list[index]

        for i in range(cycle_start_position, rows):
            if (self.main_table[i][-1] == '2' and list_append_flag != -1):
                list_append_flag = 1

            if (list_append_flag == 1):
                self.iterate_list.append(self.main_table[i])
            
            if (i==cycle_end_position):
                list_append_flag = -1
                break

    def find_inner_loop_count(self, start_row_index, input_temperature_value):
        count_seperate_flag = 0
        loop_count_value = 0

        ref_temp_col = 0
        ref_current_col = 0

        match_string = ''

        for col_index in range(0,len(self.tables[self.ref_table_number][0])):
            if (self.tables[self.ref_table_number][0][col_index].replace(" ", "").lower() == self.temperature_column_string):
                ref_temp_col = col_index

            #Check for for different discharge and charge cycles - nested loops
            if (self.tables[self.ref_table_number][0][col_index].replace(" ", "").lower() == self.ch_dch_column_string):
                count_seperate_flag = 1
                if ('Discharge' in self.main_table[start_row_index][self.comment_iterator]):
                    match_string = 'dch'
                    ref_current_col = col_index
                if ('Charge' in self.main_table[start_row_index][self.comment_iterator]):
                    match_string = 'ch'
                    ref_current_col = col_index

        #Check for counting the rows for nested loops
        if (count_seperate_flag == 1 and input_temperature_value != 'NA'):
            for row_index in range(1, len(self.tables[self.ref_table_number])):
                #Block for segregation of the steps based on temperature
                if (self.tables[self.ref_table_number][row_index][ref_current_col].replace(" ", "").lower() == match_string and self.tables[self.ref_table_number][row_index][ref_temp_col].replace(" ", "").lower() == input_temperature_value):
                    loop_count_value += 1
        else:
            #Count all the rows for normal loops
            loop_count_value = len(self.tables[self.ref_table_number]) - 1

        print('inner loop---------------------')
        print(input_temperature_value)
        print(match_string)
        print(loop_count_value)

        return loop_count_value

    def expand_the_table_steps(self, input_file, input_temperature_value):
        print('Expanding the steps\n')
        
        #Find the cycle start and end positions
        self.find_cycle_positions(len(input_file))

        print()
        
        #Adding the breaking points in the main table
        self.add_flag_bit(len(input_file))
        
        inner_loop_count = 0
        temp = []
        final_temp = []

        for x in range(0, len(input_file)):
            #Expanding steps for multiple cycles

            if (('CYCLE-START' in input_file[x][self.command_iterator]) and (len(input_file[x]) == len(input_file[0])+1)):
                #Finding the iterate list
                self.find_iterate_list(x, len(input_file))

                inner_loop_count = self.find_inner_loop_count(x, input_temperature_value)

                for i in range (0, inner_loop_count):
                    #Adding cycle loop steps
                    for x1 in self.iterate_list:
                        val = x1[:]
                        temp.append(val)   

            #Skip the previous expanded cycle
            if (len(input_file[x]) == len(input_file[0])+1):
                continue
            '''
            #Expanding steps for multiple values for a single step
            pulse_pattern = r'I=I\d+toI\d+'
            if (re.search(pulse_pattern, input_file[x][self.parameter_iterator].replace('*', ''))):
                #Derive the number of times the step has to be repeated
                find_number = int(input_file[x][self.parameter_iterator].replace('*', '')[-1])
                
                for i in range(0, find_number):
                    val = input_file[x][:]
                    temp.append(val)
                    temp.append(['', 'Rest', '', 't>t_rest_pulse', 'pulse relaxation'])
                continue
            '''
            #Appending the remaining values (which don't need to be expanded)
            val = input_file[x][:]
            temp.append(val)

        #Appending all the expanded values to the final list
        final_temp = temp[:]

        return final_temp   

    def populate_unkown_values(self, input_temperature_value, ref_table_start_row):
        print('Populating unknown values\n')
        rows = len(self.main_table_expanded)
        #columns = len(main_table_expanded[1]) - 1

        ref_temperature_column = 0
        ref_current_column = 0
        ref_cycle_column = 0
        ref_start_soc_column = 0
        ref_end_soc_column = 0
        ref_soc_column = 0
        ref_pulse_column = 0

        current_row_flag = 0
        current_pulse_flag = 1
        
        temperature_flag_val = 0
        current_flag_val = ''
        count_flag_val = ''

        #Defining the start row for the supporting table
        if (ref_table_start_row == r'NA'):
            current_row_flag = 1
        else:
            current_row_flag = ref_table_start_row

        print(self.ref_table_number)
        if (len(self.tables[self.ref_table_number][0]) == 0):
            print('-----------------------------no table found')

        #Get the reference table coulmn numbers
        for col_val in range(0,len(self.tables[self.ref_table_number][0])):
            
            #Keyword search for the cycle count values in the table
            if (self.tables[self.ref_table_number][0][col_val].replace(" ", "").lower() == self.temperature_column_string):
                #print(f'Temperature Column: {col_val}')
                ref_temperature_column = col_val

            #Keyword search for the current values in the table
            if (self.tables[self.ref_table_number][0][col_val].replace(" ", "").lower() == self.current_column_string):
                #print(f'Current Column: {col_val}')
                ref_current_column = col_val

            #Keyword search for the cycle count values in the table
            if (self.tables[self.ref_table_number][0][col_val].replace(" ", "").lower() == self.cycle_count_column_string):
                #print(f'Cycle count Column: {col_val}')
                ref_cycle_column = col_val

            #Keyword search for the Start SOC values in the table
            if (self.tables[self.ref_table_number][0][col_val].replace(" ", "").lower() == self.start_soc_column_string):
                #print(f'SOC Column: {col_val}')
                ref_start_soc_column = col_val

            #Keyword search for the End SOC values in the table
            if (self.tables[self.ref_table_number][0][col_val].replace(" ", "").lower() == self.end_soc_column_string):
                #print(f'SOC Column: {col_val}')
                ref_end_soc_column = col_val

            #Keyword search for the SOC values in the table
            if (self.tables[self.ref_table_number][0][col_val].replace(" ", "").lower() == self.soc_column_string):
                #print(f'SOC Column: {col_val}')
                ref_soc_column = col_val

            #Keyword search for the pulse time values in the table
            if (self.tables[self.ref_table_number][0][col_val].replace(" ", "").lower() == self.pulse_time_column_string):
                #print(f'SOC Column: {col_val}')
                ref_pulse_column = col_val

        for i in range(0,rows):

            #Populate the temperature values to parameter
            if ('SetTemperature' in self.main_table_expanded[i][self.command_iterator] and ('T=TSet' in self.main_table_expanded[i][self.parameter_iterator] or 'T=xx' in self.main_table_expanded[i][self.parameter_iterator])):
                #Block for single test plan
                if (input_temperature_value == r'NA'):
                    #Storing the value in temperature parameter format
                    temperature_flag_val =  str('T=' + self.tables[self.ref_table_number][current_row_flag][ref_temperature_column] + '°C')
                    #print(temperature_flag_val)
                    #Updating the temperature value from the supporting table
                    self.main_table_expanded[i][self.parameter_iterator] = temperature_flag_val
                #Block for multiple test plan
                else:
                    self.main_table_expanded[i][self.parameter_iterator] = f'T={input_temperature_value}°C'

            #Populate the current values to parameter
            if ('I=xx' in self.main_table_expanded[i][self.parameter_iterator]):
                #Storing the value in current parameter format
                current_flag_val =  str('I=' + self.tables[self.ref_table_number][current_row_flag][ref_current_column])
                if ('I=Imax' in current_flag_val):
                    self.integrate_operating_window_flag = 1
                #print(current_flag_val)
                #Updating the temperature value from the supporting table
                self.main_table_expanded[i][self.parameter_iterator] = current_flag_val

            #Populate the cycle count to parameter
            if ('CYCLE-END' in self.main_table_expanded[i][self.command_iterator] and ('COUNT=n' in self.main_table_expanded[i][self.parameter_iterator] or 'COUNT=xx' in self.main_table_expanded[i][self.parameter_iterator])):
                count_flag_val = str('COUNT=' + self.tables[self.ref_table_number][current_row_flag][ref_cycle_column])
                #print(count_flag_val)
                self.main_table_expanded[i][self.parameter_iterator] = count_flag_val

            #Populate the pulse time value  and ('Discharge' in self.main_table_expanded[i][self.command_iterator] or 'Charge' in self.main_table_expanded[i][self.command_iterator])
            if ('t>tpulse' in self.main_table_expanded[i][self.exit_condition_iterator].replace(" ","").lower() and 'pulse' in self.main_table_expanded[i][self.comment_iterator].lower() ):
                #Storing the value in temperature parameter format
                pulse_time_flag_val =  str('tpulse>' + self.tables[self.ref_table_number][current_row_flag][ref_pulse_column] + 's')
                #Updating the pulse time value from the supporting table
                self.main_table_expanded[i][self.exit_condition_iterator] = pulse_time_flag_val

            #Populate the SOCrpt value to the exit condition
            if ('soc' in self.main_table_expanded[i][self.comment_iterator].replace(" ","").lower() or 'x' in self.main_table_expanded[i][self.exit_condition_iterator].replace(" ","").lower()):
                #Charge block
                if ('ah-set>-xcrpt' in self.main_table_expanded[i][self.exit_condition_iterator].replace(" ","").lower() and 'Charge' in self.main_table_expanded[i][self.command_iterator]):
                    #Storing the SOCrpt value in the SOC exit condition format
                    charge_soc_flag_val = str('SOC>' + self.tables[self.ref_table_number][current_row_flag][ref_soc_column])
                    self.main_table_expanded[i][self.exit_condition_iterator] = charge_soc_flag_val
                #Discharge block
                if ('ah-set<-xcrpt' in self.main_table_expanded[i][self.exit_condition_iterator].replace(" ","").lower() and 'Discharge' in self.main_table_expanded[i][self.command_iterator]):
                    #Storing the SOCrpt value in the SOC exit condition format
                    discharge_soc_flag_val = str('SOC<' + self.tables[self.ref_table_number][current_row_flag][ref_soc_column])
                    self.main_table_expanded[i][self.exit_condition_iterator] = discharge_soc_flag_val
            
            #Populate the start SOC value
            if ('ah-set>-(1-x).crpt' in self.main_table_expanded[i][self.exit_condition_iterator].replace(" ","").lower() and 'startsoc' in self.main_table_expanded[i][self.comment_iterator].replace(" ","").lower()):
                charge_start_soc_flag = str('SOC>' + self.tables[self.ref_table_number][current_row_flag][ref_start_soc_column])
                self.main_table_expanded[i][self.exit_condition_iterator] = charge_start_soc_flag

            #Populate the end SOC value
            if ('ah-set>-(1-x).crpt' in self.main_table_expanded[i][self.exit_condition_iterator].replace(" ","").lower() and 'endsoc' in self.main_table_expanded[i][self.comment_iterator].replace(" ","").lower()):
                charge_end_soc_flag = str(self.tables[self.ref_table_number][current_row_flag][ref_start_soc_column] + '<SOC<' + self.tables[self.ref_table_number][current_row_flag][ref_end_soc_column])
                self.main_table_expanded[i][self.exit_condition_iterator] = charge_end_soc_flag
            
            '''

            #Populate the current value
            pulse_pattern = r'I\s*=\s*I\d+\s+to\s+I\d+'
            if (re.search(pulse_pattern, self.main_table_expanded[i][self.parameter_iterator].replace('*', ''))):
                
                #Resetting the flag after it crosses the boundary
                if (current_pulse_flag == len(self.tables[self.ref_table_number][0])):
                    current_pulse_flag = 1

                #Get the refernece table column number
                for j3 in range(0,len(self.tables[self.ref_table_number])):
                    #Keyword search for the current values in the table
                    if (self.tables[self.ref_table_number][j3][0].replace(" ", "") == 'Current'):
                        ref_pulse_row  = j3               
                
                temp_string = str('I=' + self.tables[self.ref_table_number][ref_pulse_row][current_pulse_flag].replace("-", ""))
                
                self.main_table_expanded[i][self.parameter_iterator] = temp_string
                current_pulse_flag += 1

            '''
            
            #Incrementing the supporting table row iterator value
            #Check main table cycle end with the current processing row
            if ('CYCLE-END' in self.main_table_expanded[i][self.command_iterator] and len(self.main_table_expanded[i]) == len(self.main_table_expanded[0])+1):
                current_row_flag += 1
                #Resetting the flag after it crosses the boundary (To avoid list out of range error)
                if (current_row_flag == len(self.tables[self.ref_table_number])):
                    current_row_flag = 1

    def apply_operating_window_limits(self, input_file):
        print('Integrating operating window limits\n')

        #Load the excel file
        sheet_data = pd.read_excel(self.operating_window_file_path, sheet_name=self.cell_type_input)

        discharge_current_string = 'Itherm,cont,dch[A]'.lower()
        charge_current_string = 'Iplat,cont[A]'.lower()

        temperature_col_index = 0

        temp = []
        final_temp = []

        charge_template = ['', 'Charge', 'enter_parameter_here', 'enter_exit_condition_here', 'comment']
        discharge_template = ['', 'Discharge', 'enter_parameter_here', 'enter_exit_condition_here', 'comment']

        for x in range(0, len(input_file)):
            #Check for the temperature parameter
            if('SetTemperature' in input_file[x][self.command_iterator]):
                #Extract the number from the temperature string
                find_number = re.findall(r'-?\d+\.?\d*', input_file[x][self.parameter_iterator])
                #Checking with the list and finding the column index
                for temp_value in self.all_temperature_values_with_col:
                    if(find_number[0] == str(int(temp_value[0]))):
                        temperature_col_index = temp_value[1]
            
            #Check for the Imax current and replace with the operating limit window
            if ('i=imax' in input_file[x][self.parameter_iterator].replace(" ","").lower()):
                #Charge block
                if ('Charge' in input_file[x][self.command_iterator]):
                    charge_loop_iterator_start_flag = 0
                    charge_loop_iterator_end_flag = 0
                    
                    for row_index, row in sheet_data.iterrows():
                        for col_index, value in enumerate(row):
                            #Check for 'Iplat,cont[A]' string in the excel
                            if (pd.notna(value) and charge_current_string in str(value).lower().replace(" ","")):
                                #print(sheet_data.iloc[row_index-1:,temperature_col_index])

                                for val1, val2 in zip(sheet_data.iloc[row_index:,temperature_col_index], sheet_data.iloc[row_index:,col_index-1]):
                                    #Charge loop start condition
                                    if (val2 == 5 and charge_loop_iterator_end_flag == 0):
                                        charge_loop_iterator_start_flag= 1
                                    #Charge loop end condition
                                    if ((pd.isna(val2) or val1==0) and charge_loop_iterator_start_flag == 1):
                                        charge_loop_iterator_end_flag = -1
                                    #Fetching the charge steps from operating window
                                    if (charge_loop_iterator_start_flag==1 and charge_loop_iterator_end_flag==0):
                                        charge_template[2] = f'I={val1}A'
                                        charge_template[3] = f'SOC>{val2}'
                                        charge_template[4] = 'Charging current splitted according to Operating window'
                                        #Appending the charging steps
                                        temp.append(charge_template[:])
                #Discharge block
                if ('Discharge' in input_file[x][self.command_iterator]):
                    discharge_loop_iterator_start_flag = 0
                    discharge_loop_iterator_end_flag = 0
                    discharge_current_steps_list = []

                    for row_index, row in sheet_data.iterrows():
                        for col_index, value in enumerate(row):
                            #Check for 'Itherm,cont,dch[A]' string in the excel
                            if (pd.notna(value) and discharge_current_string in str(value).lower().replace(" ","")):
                                for val1, val2 in zip(sheet_data.iloc[row_index:,temperature_col_index], sheet_data.iloc[row_index:,col_index-1]):
                                    #Discharge loop start condition
                                    if ((val2 == 0 and val1!=0) and discharge_loop_iterator_end_flag == 0):
                                        discharge_loop_iterator_start_flag = 1
                                    #Discharge loop end condition
                                    if (val2 == 100 and discharge_loop_iterator_start_flag == 1):
                                        discharge_loop_iterator_end_flag = -1
                                    #Fetching the discharge steps from the operating window
                                    if (discharge_loop_iterator_start_flag==1 and discharge_loop_iterator_end_flag==0):
                                        discharge_template[2] = f'I=-{val1}A'
                                        discharge_template[3] = f'SOC<{val2}'
                                        discharge_template[4] = 'Discharging current splitted according to Operating window'
                                        discharge_current_steps_list.append(discharge_template[:])
                    
                    #Appending the discharge steps               
                    if (discharge_loop_iterator_end_flag == -1):
                        discharge_current_steps_list.reverse() #Reversing the list for (100% SOC to 0% SOC)
                        for discharge_step_val in discharge_current_steps_list:
                            temp.append(discharge_step_val[:])
                        discharge_loop_iterator_end_flag = -2

            #Check for quick charge Imax current and replace with start and end soc
            if ('i=imax' in input_file[x][self.parameter_iterator].replace(" ","").lower() and 'quickcharge' in input_file[x][self.comment_iterator].replace(" ","").lower()):
                quick_charge_loop_iterator_start_flag = 0
                quick_charge_loop_iterator_end_flag = 0

                exit_condition_string = input_file[x][self.exit_condition_iterator]
                
                start_end_soc_values = re.findall(r'\d+', exit_condition_string)
                
                for row_index, row in sheet_data.iterrows():
                    for col_index, value in enumerate(row):
                        #Check for 'Iplat,cont[A]' string in the excel
                        if (pd.notna(value) and charge_current_string in str(value).lower().replace(" ","")):
                            #print(sheet_data.iloc[row_index-1:,temperature_col_index])
                            for val1, val2 in zip(sheet_data.iloc[row_index:,temperature_col_index], sheet_data.iloc[row_index:,col_index-1]):
                                
                                if (quick_charge_loop_iterator_start_flag==1 and quick_charge_loop_iterator_end_flag==0):
                                    charge_template[2] = f'I={val1}A'
                                    charge_template[3] = f'SOC>{val2}'
                                    charge_template[4] = 'Charging current splitted according to Operating window'
                                    temp.append(charge_template[:])

                                if (val2 == int(start_end_soc_values[0]) and quick_charge_loop_iterator_end_flag == 0):
                                    quick_charge_loop_iterator_start_flag= 1
                                
                                if (val2 == int(start_end_soc_values[1]) and quick_charge_loop_iterator_start_flag == 1):
                                    quick_charge_loop_iterator_end_flag = -1

            #Skipping the Imax current step and appending the rest
            if ('i=imax' not in input_file[x][self.parameter_iterator].replace(" ","").lower()):
                val = input_file[x][:]
                temp.append(val)

        final_temp = temp[:]

        self.main_table_expanded = final_temp[:]
            
    def set_software_columns(self):
        print('Setting software columns\n')

        column_iterator = 0
        
        #Column 1
        self.values_list = self.values_list + r'column_iterator' + " " + r'row_iterator'
        self.values_list = self.values_list + "\n"
        column_iterator +=1
        
        #Column 2
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Level'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 3
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Label'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 4
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Command'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 5
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Parameter'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 6
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Termination'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 7
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Action'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 8
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Registration'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 9
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + r'Comment'
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        self.row_iterator += 1
        
    def populate_the_values(self, label, command, parameter, termination, registration, comment):
        
        column_iterator = 0
        
        #Column 1
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + str(self.row_iterator)
        self.values_list = self.values_list + "\n"
        column_iterator +=1
        
        #Column 2
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 3
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + str(label)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 4
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + str(command)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 5
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + str(parameter)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 6
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + str(termination)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 7
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 8
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + str(registration)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
        #Column 9
        self.values_list = self.values_list + str(column_iterator) + " " + str(self.row_iterator) + " " + str(comment)
        self.values_list = self.values_list + "\n"
        column_iterator += 1
        
    def set_global_safety_limit_parameters(self):
        print('Setting global safety limit parameters\n')

        label = ''
        command = 'Start'
        parameter = ''
        termination =''
        registration = ''
        comment = ''

        #Setting global voltage safety limit parameter
        termination += f'U>{round(self.upper_voltage_limit_value + 0.1, 2)}V&t>0.5sþýU<{round(self.lower_voltage_limit_value - 0.1, 2)}V&t>0.5sþý'

        #Setting global current safety limit parameter
        termination += f'I>{self.charge_current_limit_value + 20}A&t>0.5sþýI<-{self.discharge_current_limit_value + 20}A&t>0.5sþý'

        #Setting global temperature safety limit parameter
        termination += f'T1>{self.upper_temperature_limit_value}°C&t>0.5s'
            
        self.populate_the_values(label,command,parameter,termination,registration,comment)
        self.row_iterator +=1
    
    def set_calculation_parameters(self):
        #Setting the pre-start parameters based on the test
        print('Setting calculation parameters\n')
        
        cal_sheet_name = "Additional Parameters"
        variable_column_header = "Variable"
        value_column_header = "Value"
        
        self.populate_the_values('', 'CalcOnce', f'U_dyn_max_pos_T={self.upper_voltage_limit_value}V', '', '', '')
        self.row_iterator += 1

        #Incrmenting 0.05 for the upper safety limit
        self.populate_the_values('', 'CalcOnce', f'U_dyn_max_safety_pos_T={round(self.upper_voltage_limit_value + 0.05, 2)}V', '', '', '')
        self.row_iterator += 1
        
        self.populate_the_values('', 'CalcOnce', f'U_dyn_min_pos_T={self.lower_voltage_limit_value}V', '', '', '')
        self.row_iterator += 1
        
        #Decrementing 0.05 for the lower safety limit
        self.populate_the_values('', 'CalcOnce', f'U_dyn_min_safety_pos_T={round(self.lower_voltage_limit_value - 0.05, 2)}V', '', '', '')
        self.row_iterator += 1
    
        self.populate_the_values('', 'CalcOnce', 't_relax_electric=1800s', '', '', 'Electrical relaxation time')
        self.row_iterator += 1    
        
        self.populate_the_values('', 'CalcOnce', 't_relax_thermal=3600s', '', '', 'Thermal relaxation time')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_C20_DCH=-CA/20', '', '', 'Discharge current with C/20') 
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_C10_DCH=-CA/10', '', '', 'Discharge current with C/10')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_C5_DCH=-CA/5', '', '', 'Discharge current with C/5')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_C3_DCH=-CA/3', '', '', 'Discharge current with C/3')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_C2_DCH=-CA/2', '', '', 'Discharge current with C/2')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_1C_DCH=-CA', '', '', 'Discharge current with 1C')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_2C_DCH=-2CA', '', '', 'Discharge current with 2C')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_C20_CH=CA/20', '', '', 'Charge current with C/20')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_C3_CH=CA/3', '', '', 'Charge current with C/3')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', 'I_1C_CH=CA', '', '', 'Charge current with 1C')
        self.row_iterator += 1

        self.populate_the_values('', 'CalcOnce', f'I_derating={self.derating_current_limit}A', '', '', 'De-rating current for CV charge')
        self.row_iterator += 1

        #Load the excel file and process registration parameters
        initialization_sheet_data = pd.read_excel(self.initialization_parameters_file_path, sheet_name=cal_sheet_name)

        #Populate the Calculate parameters
        for variable, value in zip(initialization_sheet_data[variable_column_header], initialization_sheet_data[value_column_header]):
            if (variable or value):
                self.populate_the_values('', 'Calculate', f'{variable}={value}', '', '', '')
                self.row_iterator += 1

        self.populate_the_values('', r'Pause', '', 't>t_relax_electric', '', '')
        self.row_iterator += 1
        
        self.populate_the_values('', r'Define', 'type=0þýURange=1', '', '', '')
        self.row_iterator += 1

    def set_registration_paramters(self):
        print('Setting registration parameters\n')
        
        registration_parameters_count = len(self.registration_parameters_input)

        registration_parameters_list = ''
        temp_list = ''

        for value in self.registration_parameters_input:
            temp_list += str(value)
            temp_list += "\n"

        #Adding Main and the number of registration parameters
        registration_parameters_list += str("Main") + "\n" + str(registration_parameters_count) + "\n"

        #Adding the registration parameters to the list
        registration_parameters_list += temp_list

        self.registration_values_list += registration_parameters_list
        
        self.set_software_columns()
        
        self.set_global_safety_limit_parameters()

    def set_cycle_start(self, iterator_val):
        #print(f"{iterator_val}: Cycle Start")

        if (self.cycler_input == 'BaSyTec'):

            if (self.main_table_expanded[iterator_val][1] == r'CYCLE-START-1'):
                label = r'PRE_CONDITION_START'
                self.charge_nom_flag = 1
            elif (self.main_table_expanded[iterator_val][1] == r'CYCLE-START-2'):
                label = r'RPT_KAPA_START'
                self.rpt_kapa_flag = 1
            elif ('CYCLE-START' in self.main_table_expanded[iterator_val][1] and len(self.main_table_expanded[iterator_val]) == 6):
                label = 'LOOP_' + str(self.cycle_loop_start_flag) + '_START'
                self.cycle_loop_start_flag += 1
            else:
                label = ''
                
            command = r'Cycle-start'
            registration = ''
            parameter = self.main_table_expanded[iterator_val][self.parameter_iterator]
            comment = self.main_table_expanded[iterator_val][self.comment_iterator]
            termination = self.main_table_expanded[iterator_val][self.exit_condition_iterator]
        
            comment = ''

            self.populate_the_values(label,command,parameter,termination,registration,comment)

    def set_cycle_end(self, iterator_val):
        #print(f"{iterator_val}: Cycle End")
        
        if (self.cycler_input == 'BaSyTec'):
            
            if (self.main_table_expanded[iterator_val][1] == 'CYCLE-END-1'):
                label = r'PRE_CONDITION_END'
            elif (self.main_table_expanded[iterator_val][1] == 'CYCLE-END-2'):
                label = r'RPT_KAPA_END'
            elif ('CYCLE-END' in self.main_table_expanded[iterator_val][1] and len(self.main_table_expanded[iterator_val]) == 6):
                label = 'LOOP_' + str(self.cycle_loop_end_flag) + '_END'
                self.cycle_loop_end_flag += 1
            else:
                label = ''

            command = r'Cycle-end'
            registration = ''
            parameter = self.main_table_expanded[iterator_val][self.parameter_iterator]
            parameter = parameter.lower()
            comment = self.main_table_expanded[iterator_val][self.comment_iterator]
            termination = self.main_table_expanded[iterator_val][self.exit_condition_iterator]   
            comment = ''
            
            self.populate_the_values(label,command,parameter,termination,registration,comment)

    def set_charge(self, iterator_val):
        #print(f"{iterator_val}: Charge") 

        if (self.cycler_input == 'BaSyTec'):

            label = ''
            command = r'Charge'
            registration = r't=3sþýU=5mV'
            parameter = self.main_table_expanded[iterator_val][self.parameter_iterator]
            comment = self.main_table_expanded[iterator_val][self.comment_iterator]
            termination = self.main_table_expanded[iterator_val][self.exit_condition_iterator]
            
            #Setting CHARGE_NOM label
            if (self.charge_nom_flag == 1):
                label = r'CHARGENOM'
                self.charge_nom_flag = 0
            
            if (parameter == 'I=IRPT/3'):
                parameter = r'I=I_C3_CHþýU=U_dyn_max_pos_T'
                
            if (parameter == 'V=Vdyn,max'):
                parameter = r'I=I_deratingþýU=U_dyn_max_pos_T'

            if (parameter == 'I=IRPT/3,V=Vdyn,max'):
                parameter = r'I=I_C3_CHþýU=U_dyn_max_pos_T'    

            if (parameter == 'I=0.05.IRPT'):
                parameter = r'I=I_RPT_C20_CH'
            
            #Initialize parameters for different C-rate dynamically
            charge_pattern_1 = r'I\s*=\s*[\d.]+[cC]' #Defining the pattern for I=xC current
            
            charge_pattern_2 = r'I\s*=\s*[cC]/[\d.]+' #Defining the pattern for I=C/x current

            #Block for setting parameter for whole number C-rates
            if (re.search(charge_pattern_1, parameter)):
                find_number = re.findall(r'\d+\.?\d*', parameter)
                parameter = f'I=I_{find_number[0]}C_CHþýU=U_dyn_max_pos_T'

            #Block for setting parameter for fraction number C-rates
            if (re.search(charge_pattern_2, parameter)):
                find_number = re.findall(r'\d+\.?\d*', parameter)
                parameter = f'I=I_C{find_number[0]}_CHþýU=U_dyn_max_pos_T'

            #Initialize pulse parameters for different currents dynamically
            pulse_pattern = r'I\s*=\s*I[\d]' #Defining the pattern for I=Ix current
            if (re.search(pulse_pattern, parameter)):
                find_number = re.findall(r'\d+\.?\d*', parameter)
                parameter = f'I=I_{find_number[0]}_CH'
                
            #Dynamically set SOC value for termination
            soc_pattern = r'SOC\s*>\s*[\d]' #Defining the pattern for SOC>x%
            if (re.search(soc_pattern, termination)):
                find_number = re.findall(r'\d+\.?\d*', termination)
                if (find_number[0]=='100.0' or find_number[0]=='100'):
                    termination = f'U>U_dyn_max_safety_pos_T&t>0.5sþýI<I_C20_CH'
                else:
                    termination = f'SOC>{find_number[0]}þýU>U_dyn_max_safety_pos_T&t>0.5sþýI<I_C20_CH'

            #Need condition for normal and nominal SOC   
            if (termination == 'V>Vdyn,max' and label=='CHARGENOM'):
                termination = r'U>U_dyn_max_pos_TþýNOM_SOC>95'
            
            if (termination == 'V>Vdyn,max'):
                termination = r'U>U_dyn_max_pos_TþýSOC>95'
                
            if (termination == 'I<0.05.IRPT'):
                termination = r'U>U_dyn_max_safety_pos_TþýI<I_C20_CH'
            
            #Adjust SOC by charge termination   
            if (termination == 'Ah-Set>-(1-SOCset)CRPT,I<0.05⋅IRPT'):
                termination = r'U>U_dyn_max_safety_pos_T&t>1sþýAh-Set>SOC_Step_CH_5þýSOC>20'
            
            #Pulse charge termination
            if (termination == 't-step>tpulse,V>Vdyn,max,Ah-Set>0(SOC>100)'):
                termination = r't>t_pulsetimeþýU>U_dyn_max_pos_TþýSOC>100'
            
            #Dynamically set pulse time for termination
            pulse_time_pattern = r'tpulse>(\d+)s' #Defining the pattern for t>x
            if (re.search(pulse_time_pattern, termination)):
                find_number = re.findall(r'\d+\.?\d*', termination)
                termination = f't>{find_number[0]}sþýU<U_dyn_min_safety_pos_T&t>0.5s'

            if (parameter == 'I=I_RPT_C20_CH'):
                termination = r'U>U_dyn_max_pos_T'
                
            if ('QC' in parameter):
                registration = r't=3sþýU=1mV'
            
            comment = ''
            
            self.populate_the_values(label,command,parameter,termination,registration,comment) 

    def set_discharge(self, iterator_val):
        #print(f"{iterator_val}: Discharge")

        if (self.cycler_input == 'BaSyTec'):

            label = ''
            command = r'Discharge'
            registration = r't=3sþýU=5mV'
            parameter = self.main_table_expanded[iterator_val][self.parameter_iterator]
            comment = self.main_table_expanded[iterator_val][self.comment_iterator].replace(" ", "").lower()
            termination = self.main_table_expanded[iterator_val][self.exit_condition_iterator]    
            
            #Setting the RPT_KAPA label for the first discharge cycle
            if (self.rpt_kapa_flag == 1):
                label = r'RPT_KAPA'
                self.rpt_kapa_flag = 0
            
            #Setting the parameters value from the test specifications
            if (parameter == 'I=INOM/3'):
                parameter = r'I=I_C3_DCHþýU=U_dyn_min_pos_T'
            
            if (parameter == 'I=INOM/3,V=Vdyn,min'):
                parameter = r'I=I_C3_DCHþýU=U_dyn_min_pos_T'

            if (parameter == 'I=0.05.IRPT'):
                parameter = r'I=I_C20_DCHþýU=U_dyn_min_pos_T'
            
            #Initialize parameters for different C-rate dynamically
            discharge_pattern_1 = r'I\s*=\s*[\d.]+[cC]' #Defining the pattern for I=xC current
            
            discharge_pattern_2 = r'I\s*=\s*[cC]/[\d.]+' #Defining the pattern for I=C/x current

            #Block for setting parameter for whole number C-rates
            if (re.search(discharge_pattern_1, parameter)):
                find_number = re.findall(r'\d+\.?\d*', parameter)
                parameter = f'I=I_{find_number[0]}C_DCHþýIpos=0A'
                self.discharge_current_list.append(f'I_{find_number[0]}C_DCH') #For defining the initialization parameters dynamically

            #Block for setting parameter for fraction number C-rates
            if (re.search(discharge_pattern_2, parameter)):
                find_number = re.findall(r'\d+\.?\d*', parameter)
                parameter = f'I=I_C{find_number[0]}_DCHþýIpos=0A'
                self.discharge_current_list.append(f'I_C{find_number[0]}_DCH') #For defining the initialization parameters dynamically
            
            #Initialize pulse parameters for different currents dynamically
            pulse_pattern = r'I\s*=\s*I[\d]' #Defining the pattern for I=Ix current
            if (re.search(pulse_pattern, parameter)):
                find_number = re.findall(r'\d+\.?\d*', parameter)
                parameter = f'I=I_{find_number[0]}_DCH'
                
            #Setting the termination value from the test specifications
            if (termination == 'V<Vdyn,min' and 'Ipos=0A' not in parameter):
                termination = r'U<U_dyn_min_safety_pos_TþýI<I_C20_DCH&t>1s'

            #Dynamically set SOC value for termination
            soc_pattern = r'SOC\s*[=<]\s*[\d]' #Defining the pattern for SOC<x% or SOC=x%
            if (re.search(soc_pattern, termination)):
                find_number = re.findall(r'\d+\.?\d*', termination)
                if (find_number[0]=='0.0' or find_number[0]=='0'):
                    termination = f'U<U_dyn_min_safety_pos_T&t>0.5sþýI<I_C20_DCH&t>1s'
                else:
                    termination = f'SOC<{find_number[0]}þýU<U_dyn_min_safety_pos_T&t>0.5sþýI<I_C20_DCH&t>1s'
                
            if (termination == 'V<Vdyn,min' and 'Ipos=0A' in parameter):
                termination = r'U<U_dyn_min_pos_T&t>1s'

            #Dynamically set SOC termination condition
            termination_pattern = r'Ah-Set<-([-+]?\d*\.?\d+)CRPT,I<([-+]?\d*\.?\d+).INOM' #Defining the pattern for 'Ah-Set<-0.5CRPT,I<0.05.INOM'
            if (re.search(termination_pattern, termination)):
                find_number = re.findall(r'\d+\.?\d*', termination)
                termination = f'SOC<{float(find_number[0])*100}þýU<U_dyn_min_safety_pos_TþýI<I_C20_DCH&t>1s'

            #Adjust SOC by discharge termination     
            if (termination == 'Ah-Set<-(1-SOCset)CRPT,I<0.05⋅INOM'):
                termination = r'U<U_dyn_min_pos_TþýAh-Set<SOC_Step_DCHþýSOC<20&t>1s'
            
            #Pulse discharge termination
            if (termination == 't-step>tpulse,V<Vdyn,min,Ah-Counter<-CRPT(SOC<0)'):
                termination = r't>t_pulsetimeþýU<U_dyn_min_pos_TþýSOC<0'
        
            #Dynamically set pulse time for termination
            pulse_time_pattern = r'tpulse>(\d+)s' #Defining the pattern for t>x
            if (re.search(pulse_time_pattern, termination)):
                find_number = re.findall(r'\d+\.?\d*', termination)
                termination = f't>{find_number[0]}sþýU<U_dyn_min_safety_pos_T&t>0.5s'
                
            if (parameter == 'I=I_RPT_C20_DCH'):
                termination = r'U<U_dyn_min_pos_T'
            
            #Process and populate the values
            if ('cvpartofcccvdischarge' not in comment and 'no other tests are planned' not in comment):
                comment = ''
                self.populate_the_values(label,command,parameter,termination,registration,comment)
            #Skip the second discharge step and termination step
            else:
                self.row_iterator += -1
        
    def rest_time(self, iterator_val):
        #print(f"{iterator_val}: Rest")
        #Lazy module imports
        import re

        if (self.cycler_input == 'BaSyTec'):

            label = ''
            command = r'Pause'
            registration = r't=10sþýU=5mV'
            parameter = self.main_table_expanded[iterator_val][self.parameter_iterator]
            comment = self.main_table_expanded[iterator_val][self.comment_iterator]
               
            if (r'pulse relaxation' in comment):
                termination = r't>t_rest_pulse'
            elif (r'thermal' in comment):
                termination = r't>t_relax_thermal'
            else:
                termination = self.main_table_expanded[iterator_val][self.exit_condition_iterator].replace('*', '')

                minute_pattern = r't>(\d+)min'
                hour_pattern = r't>(\d+)h'

                if (re.search(minute_pattern, termination)):
                    find_number = re.findall(r'\d+\.?\d*', termination)
                    termination = f't>{int(find_number[0])*60}s'
                
                elif (re.search(hour_pattern, termination)):
                    find_number = re.findall(r'\d+\.?\d*', termination)
                    termination = f't>{int(find_number[0])*60*60}s'
            
            comment = ''
            
            self.populate_the_values(label,command,parameter,termination,registration,comment) 

    def set_value(self, iterator_val):
        #print(f"{iterator_val}: Set")

        if (self.cycler_input == 'BaSyTec'):

            label = ''
            command = r'Set'
            registration = ''
            parameter = self.main_table_expanded[iterator_val][self.parameter_iterator]
            comment = ''
            termination = self.main_table_expanded[iterator_val][self.exit_condition_iterator]
            
            self.populate_the_values(label,command,parameter,termination,registration,comment) 

    def set_temperature(self, iterator_val):
        #print(f"{iterator_val}: Set Temperature")

        if (self.cycler_input == 'BaSyTec'):

            label = ''
            command = r'Set-Temp'
            registration = r't=10sþýU=5mV'
            parameter = self.main_table_expanded[iterator_val][self.parameter_iterator]
            comment = ''
            termination = self.main_table_expanded[iterator_val][self.exit_condition_iterator]
            
            self.populate_the_values(label,command,parameter,termination,registration,comment) 

    def set_termination(self):
        print('\nSetting Termination')

        if (self.cycler_input == 'BaSyTec'):

            label = 'TERMINATION'
            command = 'Discharge'
            parameter = 'I=I_C3_DCHþýU=U_dyn_min_pos_T'
            registration = ''
            comment = ''
            termination = 'SOC<30þýI>I_C20_DCH'
            
            self.populate_the_values(label,command,parameter,termination,registration,comment)
            self.row_iterator += 1

    def set_end(self):
        print('\nSetting end parameters')

        if (self.cycler_input == 'BaSyTec'):

            #Commented as of now
            '''
            populate_the_values('', 'Pause', '', 't>0.6s', 'FirstPoint=1msþýt=0.1sþýU=1mV', '')
            row_iterator += 1
            
            populate_the_values('', 'message', '', '', '', '')
            row_iterator += 1
            
            populate_the_values('', 'Pause', '', 't>1min', 't=1minþýU=1mV', '')
            row_iterator += 1
            '''
            
            label = 'STOP'
            command = 'Stop'
            registration = ''
            parameter = ''
            comment = ''
            termination = ''
            
            self.populate_the_values(label,command,parameter,termination,registration,comment)
            self.row_iterator +=1

    def switch_function(self, input_val,iterator_val):
            if (input_val == 0):
                self.set_cycle_start(iterator_val)
            elif (input_val == 1):
                self.set_cycle_end(iterator_val)
            elif (input_val == 2):
                self.set_charge(iterator_val)
            elif (input_val == 3):
                self.set_discharge(iterator_val)
            elif (input_val == 4):
                self.rest_time(iterator_val)
            elif (input_val == 5):
                self.set_value(iterator_val)
            elif (input_val == 6):
                self.set_temperature(iterator_val)    

    def process_main_table_expanded(self):
        rows = len(self.main_table_expanded)
        
        print('Processing...\n')
        for i in range(0,rows):
            #print(f"Row: {i} - {self.main_table_expanded[i][1]}") #Commented as of now
            
            for index,element in enumerate(self.doc_command_list):
                if ("CYCLE-START" in self.main_table_expanded[i][1]):
                    self.switch_function(0,i)
                    self.row_iterator+=1
                    break
                elif ("CYCLE-END" in self.main_table_expanded[i][1]):
                    self.switch_function(1,i)
                    self.row_iterator+=1
                    break
                elif (self.main_table_expanded[i][1]==element):
                    self.switch_function(index,i)
                    self.row_iterator+=1
                    break
        
        #Check for termination condittion and add if required
        if (self.termination_input == 1):
            self.set_termination()
    
    def update_rows_and_columns(self):
        #Adjusting the values of the rows and column
        self.values_list = self.values_list.replace("column_iterator", str(9)) #Updating the number of columns
        self.values_list = self.values_list.replace("row_iterator", str(self.row_iterator)) #Updating the number of rows

        #Appending the registration parameters to the list
        self.final_values_list = self.registration_values_list + self.values_list

    def find_test_spec_file_name(self):
        file_name_pattern = r'DTC[-_]P[-_]\d+[-_]\d+'

        find_file_prefix = re.search(file_name_pattern, self.test_spec_file_path)

        return find_file_prefix[0]

    def create_output_folder(self):
        output_file_prefix = self.find_test_spec_file_name()

        self.output_folder = self.base_folder + f'/{output_file_prefix}_TestScript'

    def create_output_file_name(self, temperature_val):
        generated_file_name = ''

        output_file_prefix = self.find_test_spec_file_name()

        #Adding the temperature value to the file name (For seperate test plans)
        if (temperature_val != 'NA'):
            output_file_prefix += f'_[{temperature_val}_degC]'

        generated_file_name = str(f'{output_file_prefix}_TestScript.pln')

        return generated_file_name

    def save_output_file(self, temperature_val):
        #Lazy module imports
        import os
        
        output_file_name = self.create_output_file_name(temperature_val)

        #Creating output folder
        if not os.path.exists(self.output_folder): 
            os.makedirs(self.output_folder) 
        
        output_file_path = os.path.join(self.output_folder, output_file_name)
        
        with open (output_file_path, 'w',encoding = 'ANSI') as file:
            file.write(self.final_values_list)
            file.close()
        
        print('\nTest plan generated')
        
        print(f'\nCheck the file - {output_file_name} in {self.output_folder} folder')

class ProcessBlockForGUI:
    test_plan_status_flag = 0 #Flag for tracking the test plan generation status

    def __init__(self) -> None:
        pass
    
    def get_test_procedure_inputs(file_path):
        #Lazy module imports
        from docx import Document

        # Load the document
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        paragraphs_list = list(paragraphs)
        
        tables_name=[]
        tables_list_for_gui = [] #list of tables for GUI

        for table in doc.tables:
            # Extract table name
            for line_index, paragraph in enumerate(paragraphs_list):
                preceding_paragraph = None
                #Find the paragraph one line above the table
                if line_index+1 < len(paragraphs_list) and paragraphs_list[line_index+1]._element.getnext() is table._element:
                    preceding_paragraph = paragraph.text.strip()    
                    tables_name.append(preceding_paragraph)
                    if (r'test procedure' in preceding_paragraph.lower()):
                        tables_list_for_gui.append(preceding_paragraph)
                    break
        
        return tables_list_for_gui

    def read_operating_window_file(file_path):
        #Lazy module imports
        from openpyxl import load_workbook

        #Loading the excel file
        excel_file = load_workbook(file_path, read_only=True)

        excel_sheets_name = []
        cell_types_list_for_gui = []
        keyword = 'cell'

        excel_sheets_name = excel_file.sheetnames

        cell_types_list_for_gui = [temp for temp in excel_sheets_name if keyword.lower() in temp.lower()]

        return cell_types_list_for_gui  
    
    def check_test_plan_status():
        return ProcessBlockForGUI.test_plan_status_flag
        
#Create an Global instance for accesing the functions and variables
self = CellTestPlanGenerator()

def main():
    CellTestPlanGenerator()

if __name__ == '__main__':
    main()


# #a = r"C:\Users\INIYANK\OneDrive - Daimler Truck\Windows 10 - Backup\Folder@Work\Work - Folders\BMS Projects\Cell Specifications\Python Code\Cell Test Plan Generator\Test Specifications\DTC-P-10-1_Validation_tests_V1.1_changed.docx"
# #a = r"C:\Users\INIYANK\OneDrive - Daimler Truck\Windows 10 - Backup\Folder@Work\Work - Folders\BMS Projects\Cell Specifications\Python Code\Cell Test Plan Generator\Test Specifications\DTC-P-6-1_Internal Resistance_V1.1_changed.docx"
# #a = r"C:\Users\INIYANK\OneDrive - Daimler Truck\Windows 10 - Backup\Folder@Work\Work - Folders\BMS Projects\Cell Specifications\Python Code\Cell Test Plan Generator\Test Specifications\DTC-P-6-2_Internal Resistances BCC_V1.1_draft.docx"
# a = r"C:\Users\INIYANK\OneDrive - Daimler Truck\Windows 10 - Backup\Folder@Work\Work - Folders\BMS Projects\Cell Specifications\Python Code\Cell Test Plan Generator\Test Specifications\DTC-P-2-2_Capacity_Energy_Efficiency_Temp_V1.0.docx"
# "C:\Users\GSMANOJ\Downloads\Cell Test Plan Generator 3\Cell Test Plan Generator\Test Specifications\DTC-P-10-1_Validation_tests_V1.1_changed.docx"
# C:\Users\GSMANOJ\Downloads\Cell Test Plan Generator_Copy\Cell Test Plan Generator\Supporting Documents\Operating Window_HDMD_CATL_B-sample_v2.8.xlsx
# C:\Users\GSMANOJ\Downloads\Cell Test Plan Generator_Copy\Cell Test Plan Generator\Test Specifications\DTC-P-2-2_Capacity_Energy_Efficiency_Temp_V1.0.docx
# C:\Users\GSMANOJ\Downloads\new\Cell Test Plan Generator_Copy\Cell Test Plan Generator\Test Specifications\DTC-P-2-1_Capacity_Energy_Efficiency_Crates_V1.0.docx
a = r"C:\Users\GSMANOJ\Downloads\Cell Test Plan Generator 3\Cell Test Plan Generator\Test Specifications\DTC-P-2-2_Capacity_Energy_Efficiency_Temp_V1.0.docx"

b = r"C:\Users\GSMANOJ\Downloads\Cell Test Plan Generator_Copy\Cell Test Plan Generator\Supporting Documents\Operating Window_HDMD_CATL_B-sample_v2.8.xlsx"

c = r"C:\Users\GSMANOJ\Downloads\Cell Test Plan Generator_Copy\Cell Test Plan Generator\Supporting Documents\Initialization_Parameters_Document.xlsx"

#d = 'Table 3 : Test procedure for thermal relaxation'
#d = 'Table 2: : Test procedure for determination of internal resistance'
# d = 'Table 5: Test procedure for Full DoD continuous current experiment'
# d = 'Table 9: Test procedure for Charge neutral pulse experiment'
d = 'Table 2: Test procedure for determination of CC/3,y, EC/3,y at different temperatures'
# d="Table 7: Test procedure for Partial DoD continuous current experiment"
# d="Table 2: Test procedure for determination of capacity CRPT, CC/x,25°C, EC/x 25°C and ηC/x,25°C at different discharging C-rates."

e = 'Cell Alpha - Op-Window'

f = 'BaSyTec'

g = 0

h = 0

i = ['Ah-Set', 'Ah-Charge', 'Ah-Discharge']


# CellTestPlanGenerator.execute_main_generate_function(a, b, c, d, e, f, g, h, i)

CellTestPlanGenerator.simulation_function(a,d,"pybamm",b,c,e)
